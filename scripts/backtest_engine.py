#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
策略回测引擎 v1.0.0
功能：执行策略回测，计算性能指标，生成.docx报告
依赖：pandas, numpy, matplotlib, python-docx
可选依赖：akshare（用于获取实际行情数据）
"""

try:
    import akshare as ak
except ImportError:
    ak = None

import pandas as pd
import numpy as np
import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import matplotlib.font_manager as fm
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings
warnings.filterwarnings('ignore')

# 设置中文字体
try:
    # 方法1：使用系统字体（Windows）
    import matplotlib.pyplot as plt
    plt.rcParams['font.sans-serif'] = ['SimHei', 'SimSun', 'Microsoft YaHei']
    plt.rcParams['axes.unicode_minus'] = False  # 解决负号显示问题
except Exception as e:
    print(f"中文字体设置失败：{e}")
    plt.rcParams['font.family'] = 'SimHei'
except:
    pass

plt.rcParams['axes.unicode_minus'] = False


def get_historical_data(symbol, start_date, end_date, asset_type='stock'):
    """
    获取历史行情数据
    
    参数：
    - symbol: 股票代码或基金代码
    - start_date: 开始日期（YYYYMMDD）
    - end_date: 结束日期（YYYYMMDD）
    - asset_type: 资产类型（'stock' 或 'fund'）
    
    返回：
    - df: 包含历史行情的DataFrame
    """
    try:
        if asset_type == 'stock':
            df = ak.stock_zh_a_hist(
                symbol=symbol,
                period="daily",
                start_date=start_date,
                end_date=end_date,
                adjust="qfq"  # 前复权
            )
            # 标准化列名
            df = df.rename(columns={
                '日期': '日期',
                '开盘': '开盘',
                '收盘': '收盘',
                '最高': '最高',
                '最低': '最低',
                '成交量': '成交量',
                '成交额': '成交额',
                '涨跌幅': '涨跌幅',
                '涨跌额': '涨跌额',
                '换手率': '换手率'
            })
        elif asset_type == 'fund':
            df = ak.fund_open_fund_info_em(
                symbol=symbol,
                indicator="单位净值走势"
            )
            # 基金数据需要特殊处理
            df['日期'] = pd.to_datetime(df['净值日期'])
            df['收盘'] = df['单位净值']
            df = df.sort_values('日期')
        
        df['日期'] = pd.to_datetime(df['日期'])
        df = df.sort_values('日期')
        
        # 计算技术指标
        df['MA5'] = df['收盘'].rolling(window=5).mean()
        df['MA20'] = df['收盘'].rolling(window=20).mean()
        
        return df
    except Exception as e:
        print(f"获取数据失败：{e}")
        return None


def run_backtest(strategy_func, df, initial_cash=1000000, 
                commission=0.0003, stamp_duty=0.001, slippage=0.002):
    """
    执行回测
    
    参数：
    - strategy_func: 策略函数（接收row和history，返回"BUY"/"SELL"/"HOLD"）
    - df: 历史行情DataFrame
    - initial_cash: 初始资金（默认100万）
    - commission: 手续费率（默认0.03%）
    - stamp_duty: 印花税率（默认0.1%，卖出时收取）
    - slippage: 滑点（默认0.2%）
    
    返回：
    - portfolio: 每日资产组合列表
    - trades: 交易记录列表
    - cash_history: 每日现金列表
    - position_history: 每日持仓数量列表
    """
    # 标准化列名（支持中文和英文列名）
    column_mapping = {
        '日期': 'date', 'date': 'date',
        '开盘': 'open', 'open': 'open',
        '最高': 'high', 'high': 'high',
        '最低': 'low', 'low': 'low',
        '收盘': 'close', 'close': 'close',
        '成交量': 'volume', 'volume': 'volume'
    }
    
    df = df.rename(columns=column_mapping)
    
    cash = initial_cash
    position = 0
    portfolio = [initial_cash]  # 第0天（未开始交易）
    cash_history = [initial_cash]
    position_history = [0]
    trades = []
    
    for i, (idx, row) in enumerate(df.iterrows()):
        # 获取历史数据（当前时点之前）
        history = df.iloc[:i] if i > 0 else df.iloc[0:0]
        
        # 获取交易信号
        try:
            signal = strategy_func(row, history)
        except Exception as e:
            print(f"策略执行错误（{row.get('date', row.get('日期', '未知'))}）：{e}")
            signal = "HOLD"
        
        # 买入
        if signal == "BUY" and cash > 0:
            buy_price = row['close'] * (1 + slippage)
            buy_amount = cash * (1 - commission)
            position = buy_amount / buy_price
            cash = 0
            trades.append({
                'date': row.get('date', row.get('日期', '未知')),
                'action': 'BUY',
                'price': buy_price,
                'amount': position,
                'cash': cash
            })
        
        # 卖出
        elif signal == "SELL" and position > 0:
            sell_price = row.get('close', row.get('收盘', 0)) * (1 - slippage)
            sell_amount = position * sell_price * (1 - commission - stamp_duty)
            cash = sell_amount
            trades.append({
                'date': row.get('date', row.get('日期', '未知')),
                'action': 'SELL',
                'price': sell_price,
                'amount': cash,
                'position': 0
            })
            position = 0
        
        # 记录当日资产
        total = cash + position * row['close']
        portfolio.append(total)
        cash_history.append(cash)
        position_history.append(position)
    
    return portfolio, trades, cash_history, position_history


def calculate_metrics(portfolio, df, risk_free_rate=0.025):
    """
    计算性能指标
    
    参数：
    - portfolio: 每日资产组合列表
    - df: 历史行情DataFrame
    - risk_free_rate: 无风险利率（默认2.5%）
    
    返回：
    - metrics: 性能指标字典
    """
    portfolio = np.array(portfolio)
    returns = np.diff(portfolio) / portfolio[:-1]
    
    # 总收益率
    total_return = (portfolio[-1] - portfolio[0]) / portfolio[0]
    
    # 年化收益率
    trading_days = len(portfolio)
    annual_return = (1 + total_return) ** (252 / trading_days) - 1 if trading_days > 0 else 0
    
    # 年化波动率
    daily_returns = pd.Series(returns)
    annual_volatility = daily_returns.std() * np.sqrt(252) if len(daily_returns) > 0 else 0
    
    # 夏普比率
    sharpe_ratio = (annual_return - risk_free_rate) / annual_volatility if annual_volatility > 0 else 0
    
    # 最大回撤
    peak = np.maximum.accumulate(portfolio)
    drawdown = (peak - portfolio) / peak
    max_drawdown = np.max(drawdown) if len(drawdown) > 0 else 0
    
    # 胜率
    winning_trades = np.sum(returns > 0)
    total_trades = np.sum(returns != 0)
    win_rate = winning_trades / total_trades if total_trades > 0 else 0
    
    # 盈亏比
    avg_win = np.mean(returns[returns > 0]) if np.sum(returns > 0) > 0 else 0
    avg_loss = np.mean(np.abs(returns[returns < 0])) if np.sum(returns < 0) > 0 else 0
    profit_loss_ratio = avg_win / avg_loss if avg_loss > 0 else np.inf
    
    # 索提诺比率
    downside_returns = returns[returns < risk_free_rate / 252]
    downside_volatility = np.std(downside_returns) * np.sqrt(252) if len(downside_returns) > 0 else 0
    sortino_ratio = (annual_return - risk_free_rate) / downside_volatility if downside_volatility > 0 else 0
    
    # 卡玛比率
    calmar_ratio = annual_return / abs(max_drawdown) if max_drawdown != 0 else np.inf
    
    metrics = {
        '总收益率': total_return,
        '年化收益率': annual_return,
        '年化波动率': annual_volatility,
        '夏普比率': sharpe_ratio,
        '最大回撤': max_drawdown,
        '胜率': win_rate,
        '盈亏比': profit_loss_ratio,
        '索提诺比率': sortino_ratio,
        '卡玛比率': calmar_ratio
    }
    
    return metrics


def plot_net_value_curve(portfolio, output_path):
    """
    绘制净值曲线图
    
    参数：
    - portfolio: 每日资产组合列表
    - output_path: 输出图片路径
    """
    plt.figure(figsize=(12, 6))
    plt.plot(range(len(portfolio)), portfolio, label='策略净值', linewidth=2, color='#1f77b4')
    
    # 添加基准线（初始资金水平线）
    plt.axhline(y=portfolio[0], color='r', linestyle='--', label='基准净值')
    
    plt.xlabel('交易日')
    plt.ylabel('净值（元）')
    plt.title('策略净值曲线')
    plt.legend()
    plt.grid(True, alpha=0.3)
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()


def plot_drawdown_curve(portfolio, output_path):
    """
    绘制回撤曲线图
    
    参数：
    - portfolio: 每日资产组合列表
    - output_path: 输出图片路径
    """
    portfolio = np.array(portfolio)
    peak = np.maximum.accumulate(portfolio)
    drawdown = (peak - portfolio) / peak * 100  # 转换为百分比
    
    plt.figure(figsize=(12, 6))
    plt.fill_between(range(len(drawdown)), drawdown, 0, color='#ff7f0e', alpha=0.3)
    plt.plot(range(len(drawdown)), drawdown, color='#ff7f0e', linewidth=1)
    
    plt.xlabel('交易日')
    plt.ylabel('回撤（%）')
    plt.title('策略回撤曲线')
    plt.grid(True, alpha=0.3)
    plt.savefig(output_path, dpi=300, bbox_inches='tight')
    plt.close()


def generate_backtest_report(strategy_name, symbol, start_date, end_date,
                           portfolio, trades, metrics, output_path,
                           net_value_img, drawdown_img):
    """
    生成回测报告（.docx格式）
    
    参数：
    - strategy_name: 策略名称
    - symbol: 股票/基金代码
    - start_date: 开始日期
    - end_date: 结束日期
    - portfolio: 每日资产组合列表
    - trades: 交易记录列表
    - metrics: 性能指标字典
    - output_path: 输出报告路径
    - net_value_img: 净值曲线图路径
    - drawdown_img: 回撤曲线图路径
    """
    doc = Document()
    
    # 页面边距
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.2)
    section.right_margin = Cm(3.2)
    
    # 封面
    p = doc.add_paragraph()
    p.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f'{strategy_name} 回测报告')
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f'标的：{symbol}')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    p = doc.add_paragraph()
    p.alignment = 1
    run = p.add_run(f'回测区间：{start_date} ~ {end_date}')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    doc.add_paragraph()
    
    # 核心结论
    doc.add_heading('核心结论（执行摘要）', level=1)
    
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'
    
    metrics_cn = [
        ('总收益率', f"{metrics['总收益率']*100:.2f}%"),
        ('年化收益率', f"{metrics['年化收益率']*100:.2f}%"),
        ('夏普比率', f"{metrics['夏普比率']:.2f}"),
        ('最大回撤', f"{metrics['最大回撤']*100:.2f}%"),
        ('胜率', f"{metrics['胜率']*100:.2f}%"),
        ('盈亏比', f"{metrics['盈亏比']:.2f}"),
        ('索提诺比率', f"{metrics['索提诺比率']:.2f}"),
        ('卡玛比率', f"{metrics['卡玛比率']:.2f}")
    ]
    
    for i, (k, v) in enumerate(metrics_cn):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v
    
    doc.add_paragraph()
    
    # 策略评级
    sharpe = metrics['夏普比率']
    if sharpe >= 2.0:
        rating = '优秀'
    elif sharpe >= 1.0:
        rating = '良好'
    elif sharpe >= 0.5:
        rating = '一般'
    else:
        rating = '较差'
    
    p = doc.add_paragraph()
    run = p.add_run(f'策略评级：{rating}')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    
    # 净值曲线图
    doc.add_heading('净值曲线', level=1)
    doc.add_picture(net_value_img, width=Inches(6))
    
    # 回撤曲线图
    doc.add_heading('回撤曲线', level=1)
    doc.add_picture(drawdown_img, width=Inches(6))
    
    # 交易明细
    if len(trades) > 0:
        doc.add_heading('交易明细', level=1)
        
        table = doc.add_table(rows=len(trades)+1, cols=4)
        table.style = 'Table Grid'
        
        # 表头
        headers = ['日期', '操作', '价格', '数量/金额']
        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h
        
        # 交易记录
        for i, trade in enumerate(trades, 1):
            table.rows[i].cells[0].text = str(trade['date'])[:10]
            table.rows[i].cells[1].text = trade['action']
            table.rows[i].cells[2].text = f"{trade['price']:.2f}"
            table.rows[i].cells[3].text = f"{trade['amount']:.2f}"
    
    # 风险提示
    doc.add_heading('风险提示', level=1)
    doc.add_paragraph('1. 数据局限性：回测使用历史数据，未来表现可能不同。')
    doc.add_paragraph('2. 模型假设限制：回测假设理想交易条件，实际交易可能存在滑点、流动性限制。')
    doc.add_paragraph('3. 过拟合风险：策略参数可能过度优化，导致样本外表现下降。')
    
    # 免责声明
    doc.add_heading('免责声明', level=1)
    doc.add_paragraph('本报告仅供参考，不构成投资建议。投资有风险，入市需谨慎。')
    doc.add_paragraph(f'数据来源：AkShare，获取时间：{pd.Timestamp.now().strftime("%Y-%m-%d %H:%M")}')
    
    # 保存
    doc.save(output_path)
    print(f'报告已生成：{output_path}')


def main():
    """
    主函数：示例用法
    """
    # 示例策略：MA均线策略
    def ma_cross_strategy(row, history):
        if len(history) < 20:
            return "HOLD"
        
        ma5 = history['收盘'].tail(5).mean()
        ma20 = history['收盘'].tail(20).mean()
        
        # 获取前一日的均线值
        if len(history) >= 21:
            prev_ma5 = history['收盘'].tail(6).head(5).mean()
            prev_ma20 = history['收盘'].tail(21).head(20).mean()
        else:
            prev_ma5 = ma5
            prev_ma20 = ma20
        
        # 金叉：MA5上穿MA20
        if prev_ma5 <= prev_ma20 and ma5 > ma20:
            return "BUY"
        # 死叉：MA5下穿MA20
        elif prev_ma5 >= prev_ma20 and ma5 < ma20:
            return "SELL"
        else:
            return "HOLD"
    
    # 参数
    strategy_name = "MA均线策略"
    symbol = "600519"
    start_date = "20240101"
    end_date = "20260526"
    asset_type = "stock"
    
    # 获取数据
    print("正在获取数据...")
    df = get_historical_data(symbol, start_date, end_date, asset_type)
    
    if df is None:
        print("获取数据失败，退出")
        return
    
    print(f"数据获取成功，共{len(df)}条记录")
    
    # 执行回测
    print("正在执行回测...")
    portfolio, trades, cash_history, position_history = run_backtest(
        ma_cross_strategy, df, initial_cash=1000000
    )
    
    # 计算性能指标
    print("正在计算性能指标...")
    metrics = calculate_metrics(portfolio, df)
    
    # 绘制图表
    print("正在绘制图表...")
    net_value_img = "net_value_curve.png"
    drawdown_img = "drawdown_curve.png"
    
    plot_net_value_curve(portfolio, net_value_img)
    plot_drawdown_curve(portfolio, drawdown_img)
    
    # 生成报告
    print("正在生成报告...")
    output_path = f"{strategy_name}_回测报告_{pd.Timestamp.now().strftime('%Y-%m-%d')}.docx"
    
    generate_backtest_report(
        strategy_name, symbol, start_date, end_date,
        portfolio, trades, metrics, output_path,
        net_value_img, drawdown_img
    )
    
    print(f"回测完成！报告已保存至：{output_path}")


if __name__ == "__main__":
    main()
